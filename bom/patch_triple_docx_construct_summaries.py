#!/usr/bin/env python3
"""Fix remaining Greek-style Lamb/Son phrases in BoM docx (chapter summaries, etc.)."""
from __future__ import annotations

import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DEFAULT_DOCX = Path(r"C:\Users\chris\Desktop\Publishing\Triple_Combination_Hebrew_FIXED.docx")

REPLACEMENTS = [
    ("בֶּן־הָאֱלֹהִים", "בֶּן־אֱלֹהִים"),
    ("בֶּן הָאֱלֹהִים", "בֶּן אֱלֹהִים"),
    ("שֵׂה הָאֱלֹהִים", "שֵׂה־אֱלֹהִים"),
    ("שֶׂה הָאֱלֹהִים", "שֶׂה־אֱלֹהִים"),
    ("בְּשֵׂה הָאֱלֹהִים", "בְּשֵׂה־אֱלֹהִים"),
    ("שֵׂה אֱלֹהִים", "שֵׂה־אֱלֹהִים"),
    ("שֶׂה אֱלֹהִים", "שֶׂה־אֱלֹהִים"),
    ("בְּשֵׂה אֱלֹהִים", "בְּשֵׂה־אֱלֹהִים"),
]

VERSE = re.compile(r"^([\u05d0-\u05ea]{1,3})\.\s*(.*)$", re.DOTALL)


def strip_nikkud(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if unicodedata.category(c) != "Mn")


def find_bom_start(paragraphs) -> int:
    for i, p in enumerate(paragraphs):
        if p.style and (p.style.name or "").startswith("Heading"):
            if strip_nikkud((p.text or "").strip()) == "נפי א":
                return i
    raise RuntimeError("BoM start not found")


def find_dc_end(paragraphs) -> int:
    for i, p in enumerate(paragraphs):
        t = strip_nikkud(p.text or "")
        if "תורה" in t and "בריתות" in t and "גלוי" in t:
            return i
    raise RuntimeError("D&C intro not found")


def set_rtl_body(p, text: str) -> None:
    for child in list(p._element):
        if child.tag.endswith("}r"):
            p._element.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    run = p.add_run(text)
    run.font.name = "David"
    run.font.size = Pt(12)


def fix_text(t: str) -> str:
    for old, new in REPLACEMENTS:
        t = t.replace(old, new)
    return t


def main() -> None:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    doc = Document(path)
    paras = doc.paragraphs
    # Full document: construct fixes apply in BoM, D&C, PGP, etc.
    n = 0
    for p in paras:
        raw = p.text or ""
        if not any(old in raw for old, _ in REPLACEMENTS):
            continue
        fixed = fix_text(raw)
        if fixed == raw:
            continue
        m = VERSE.match(raw.strip())
        if m:
            sys.path.insert(0, str(Path(__file__).resolve().parent))
            from patch_triple_docx_bom_hebrew import replace_paragraph_verse

            replace_paragraph_verse(p, m.group(1), fix_text(m.group(2).rstrip("׃").strip()))
        else:
            set_rtl_body(p, fixed)
        n += 1
    doc.save(path)
    print(f"Fixed {n} paragraphs in {path}")


if __name__ == "__main__":
    main()
