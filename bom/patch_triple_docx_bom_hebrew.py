#!/usr/bin/env python3
"""
Sync Book of Mormon Hebrew verse text in the Triple Combination docx from bom/verses/*.js.

Updates verse paragraphs (e.g. «א. …») in the BoM body only; leaves D&C/PGP untouched.
Preserves bold verse-number prefix + body run pattern used in the manuscript.

Usage:
  python bom/patch_triple_docx_bom_hebrew.py [path.docx]
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = Path(r"C:\Users\chris\Desktop\Publishing\Triple_Combination_Hebrew_FIXED.docx")
EXPORT_JSON = ROOT / "tools" / "bom_hebrew_export.json"
EXPORT_JS = ROOT / "tools" / "export_bom_hebrew.js"

VERSE_START = re.compile(r"^([\u05d0-\u05ea]{1,3})\.\s*(.*)$", re.DOTALL)
CHAPTER_HEADING = re.compile(r"^\s*פרק\s+([\u05d0-\u05ea]{1,3})\s*$")

BOOK_FROM_TITLE = {
    "נפי א": "1 Nephi",
    "נפי ב": "2 Nephi",
    "יעקב": "Jacob",
    "אנוש": "Enos",
    "ירום": "Jarom",
    "אמני": "Omni",
    "דברי מורמון": "Words of Mormon",
    "מושיה": "Mosiah",
    "אלמא": "Alma",
    "הילמן": "Helaman",
    "נפי ג": "3 Nephi",
    "נפי ד": "4 Nephi",
    "מורמון": "Mormon",
    "אתר": "Ether",
    "מורוני": "Moroni",
}


def strip_nikkud(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if unicodedata.category(c) != "Mn")


def hebrew_numeral_to_int(s: str) -> int:
    s = strip_nikkud(s.strip())
    if not s:
        raise ValueError("empty numeral")
    ones = {"א": 1, "ב": 2, "ג": 3, "ד": 4, "ה": 5, "ו": 6, "ז": 7, "ח": 8, "ט": 9}
    tens = {"י": 10, "כ": 20, "ל": 30, "מ": 40, "נ": 50, "ס": 60, "ע": 70, "פ": 80, "צ": 90}
    hundreds = {"ק": 100, "ר": 200, "ש": 300, "ת": 400}
    if s == "טו":
        return 15
    if s == "טז":
        return 16
    total = 0
    i = 0
    while i < len(s):
        if s[i : i + 2] in ("ת",) and i + 1 < len(s):
            # 500-900 as repeated ת — rare in chapter nums
            pass
        ch = s[i]
        if ch in hundreds:
            total += hundreds[ch]
        elif ch in tens:
            total += tens[ch]
        elif ch in ones:
            total += ones[ch]
        else:
            raise ValueError(f"unknown hebrew numeral char: {ch!r} in {s!r}")
        i += 1
    return total


def _set_run_font(run, name: str = "David", size_pt: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _set_paragraph_rtl(p) -> None:
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        from docx.oxml import OxmlElement

        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def replace_paragraph_verse(p, prefix: str, body: str) -> None:
    """Replace paragraph with bold «label. » + body, RTL."""
    for child in list(p._element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p._element.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_rtl(p)
    r1 = p.add_run(f"{prefix}. ")
    _set_run_font(r1, bold=True)
    text = body.rstrip()
    needs_sof = not text.endswith("׃")
    r2 = p.add_run(text + ("׃" if needs_sof else ""))
    _set_run_font(r2, bold=False)


def find_dc_intro_index(paragraphs) -> int:
    for i, p in enumerate(paragraphs):
        t = strip_nikkud(p.text or "")
        if "תורה" in t and "בריתות" in t and "היא" in t and "גלוי" in t:
            return i
    raise RuntimeError("Could not locate D&C intro paragraph.")


def find_bom_start_index(paragraphs) -> int:
    for i, p in enumerate(paragraphs):
        if not (p.style and (p.style.name or "").startswith("Heading")):
            continue
        if strip_nikkud((p.text or "").strip()) == "נפי א":
            return i
    raise RuntimeError("Could not locate BoM book title «נפי א».")


def _is_chapter_heading(p) -> bool:
    st = (p.style.name if p.style else "") or ""
    if not st.startswith("Heading"):
        return False
    return bool(CHAPTER_HEADING.match(strip_nikkud((p.text or "").strip())))


def _is_book_heading(p) -> bool:
    st = (p.style.name if p.style else "") or ""
    if st != "Heading 1":
        return False
    return strip_nikkud((p.text or "").strip()) in BOOK_FROM_TITLE


def load_export() -> dict[tuple[str, int, str], str]:
    if not EXPORT_JSON.is_file() or EXPORT_JSON.stat().st_mtime < EXPORT_JS.stat().st_mtime:
        subprocess.run(["node", str(EXPORT_JS)], cwd=str(ROOT), check=True)
    rows = json.loads(EXPORT_JSON.read_text(encoding="utf-8"))
    out: dict[tuple[str, int, str], str] = {}
    for r in rows:
        key = (r["book"], int(r["chapter"]), r["verse_label"])
        out[key] = r["hebrew"]
    return out


def patch_docx(docx_path: Path) -> tuple[int, int]:
    lookup = load_export()
    doc = Document(docx_path)
    paras = doc.paragraphs
    bom_start = find_bom_start_index(paras)
    bom_end = find_dc_intro_index(paras)

    book = None
    chapter = None
    updated = 0
    missing = 0
    lookup_misses: list[str] = []

    for p in paras[bom_start:bom_end]:
        text = (p.text or "").strip()
        if not text:
            continue

        if _is_book_heading(p):
            book = BOOK_FROM_TITLE[strip_nikkud(text)]
            chapter = None
            continue

        if _is_chapter_heading(p):
            m = CHAPTER_HEADING.match(strip_nikkud(text))
            if m:
                chapter = hebrew_numeral_to_int(m.group(1))
            continue

        m = VERSE_START.match(text)
        if not m or book is None or chapter is None:
            continue

        label = m.group(1)
        key = (book, chapter, label)
        new_heb = lookup.get(key)
        if new_heb is None:
            missing += 1
            lookup_misses.append(f"{book} {chapter}:{label} | {text[:80]}")
            continue

        old_body = m.group(2).strip().rstrip("׃").strip()
        new_body = new_heb.strip()
        if strip_nikkud(old_body) == strip_nikkud(new_body):
            continue

        replace_paragraph_verse(p, label, new_body)
        updated += 1

    backup = docx_path.with_suffix(".docx.bak")
    if not backup.exists():
        shutil.copy2(docx_path, backup)
    doc.save(docx_path)
    return updated, missing, lookup_misses


def main() -> None:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not path.is_file():
        raise SystemExit(f"File not found: {path}")
    updated, missing, lookup_misses = patch_docx(path)
    print(f"Patched {path}")
    print(f"  verses updated: {updated}")
    print(f"  lookup misses: {missing}")
    if lookup_misses:
        miss_path = ROOT / "tools" / "_docx_patch_misses.txt"
        miss_path.write_text("\n".join(lookup_misses[:200]), encoding="utf-8")
        print(f"  miss sample written: {miss_path}")


if __name__ == "__main__":
    main()
