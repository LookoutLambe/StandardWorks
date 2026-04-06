#!/usr/bin/env python3
"""
Append D&C and PGOP to the existing Hebrew BOM DOCX file,
using the exact same styles and section break patterns.
"""

import json, re, sys, copy
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

BASE = 'C:/Users/chris/Desktop/Standard Works Project'
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
nsmap = {'w': NS}


def load_dc():
    with open(f'{BASE}/dc_official_verses.json', 'r', encoding='utf-8') as f:
        return json.load(f)

def load_pgp():
    with open(f'{BASE}/pgp_official_verses.json', 'r', encoding='utf-8') as f:
        return json.load(f)

def load_chron():
    with open(f'{BASE}/dc_chronology.json', 'r', encoding='utf-8') as f:
        return json.load(f)


def add_section_break_1col(doc):
    """Add a continuous section break switching to 1 column."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = etree.SubElement(pPr, f'{{{NS}}}sectPr')
    cols = etree.SubElement(sectPr, f'{{{NS}}}cols')
    cols.set(f'{{{NS}}}num', '1')
    cols.set(f'{{{NS}}}space', '720')
    tp = etree.SubElement(sectPr, f'{{{NS}}}type')
    tp.set(f'{{{NS}}}val', 'continuous')
    return p


def add_section_break_2col(doc):
    """Add a continuous section break switching to 2 columns with separator."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = etree.SubElement(pPr, f'{{{NS}}}sectPr')
    cols = etree.SubElement(sectPr, f'{{{NS}}}cols')
    cols.set(f'{{{NS}}}num', '2')
    cols.set(f'{{{NS}}}sep', '1')
    cols.set(f'{{{NS}}}space', '144')
    tp = etree.SubElement(sectPr, f'{{{NS}}}type')
    tp.set(f'{{{NS}}}val', 'continuous')
    return p


def add_page_break_1col(doc):
    """Add a next-page section break switching to 1 column."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = etree.SubElement(pPr, f'{{{NS}}}sectPr')
    cols = etree.SubElement(sectPr, f'{{{NS}}}cols')
    cols.set(f'{{{NS}}}num', '1')
    cols.set(f'{{{NS}}}space', '720')
    tp = etree.SubElement(sectPr, f'{{{NS}}}type')
    tp.set(f'{{{NS}}}val', 'nextPage')
    return p


def add_book_title(doc, title):
    """Add a book title in Heading 1 style."""
    p = doc.add_heading(title, level=1)
    return p


def add_chapter_heading(doc, ch_num, label='פרק'):
    """Add a chapter heading in Heading 2 style."""
    from _build_dc_pgp_verses import heb_num_to_int
    # Use Hebrew numeral for chapter
    H_ONES = ['', 'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט']
    H_TENS = ['', 'י', 'כ', 'ל', 'מ', 'נ', 'ס', 'ע', 'פ', 'צ']
    H_HUNDREDS = ['', 'ק', 'ר', 'ש', 'ת']

    def int_to_heb(n):
        if n <= 0: return str(n)
        r = ''
        h = n // 100
        if h > 0:
            r += H_HUNDREDS[h] if h <= 4 else ('ת' + H_HUNDREDS[h - 4])
            n %= 100
        if n == 15: return r + 'טו'
        if n == 16: return r + 'טז'
        t = n // 10
        if t > 0: r += H_TENS[t]; n %= 10
        if n > 0: r += H_ONES[n]
        return r

    heb_ch = int_to_heb(ch_num)
    p = doc.add_heading(f'{label} {heb_ch}', level=2)
    # Ensure bold, black, centered
    for run in p.runs:
        run.font.bold = True
        run.font.color.rgb = None  # black (inherit)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set RTL
    pPr = p._p.get_or_add_pPr()
    bidi = etree.SubElement(pPr, f'{{{NS}}}bidi')
    return p


def add_verse(doc, verse_num, hebrew_text):
    """Add a verse paragraph in Normal style with David font, justified, bidi."""
    # Format: verse_num. hebrew_text
    H_ONES = ['', 'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט']
    H_TENS = ['', 'י', 'כ', 'ל', 'מ', 'נ', 'ס', 'ע', 'פ', 'צ']
    H_HUNDREDS = ['', 'ק', 'ר', 'ש', 'ת']

    def int_to_heb(n):
        if n <= 0: return str(n)
        r = ''
        h = n // 100
        if h > 0:
            r += H_HUNDREDS[h] if h <= 4 else ('ת' + H_HUNDREDS[h - 4])
            n %= 100
        if n == 15: return r + 'טו'
        if n == 16: return r + 'טז'
        t = n // 10
        if t > 0: r += H_TENS[t]; n %= 10
        if n > 0: r += H_ONES[n]
        return r

    heb_v = int_to_heb(verse_num)
    text = f'{heb_v}. {hebrew_text}'

    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.name = 'David'

    # Set RTL and justify
    pPr = p._p.get_or_add_pPr()
    bidi = etree.SubElement(pPr, f'{{{NS}}}bidi')
    jc = pPr.find(f'{{{NS}}}jc')
    if jc is None:
        jc = etree.SubElement(pPr, f'{{{NS}}}jc')
    jc.set(f'{{{NS}}}val', 'both')

    return p


def add_colophon(doc, text):
    """Add colophon text as Normal paragraph."""
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.name = 'David'
    pPr = p._p.get_or_add_pPr()
    bidi = etree.SubElement(pPr, f'{{{NS}}}bidi')
    jc = pPr.find(f'{{{NS}}}jc')
    if jc is None:
        jc = etree.SubElement(pPr, f'{{{NS}}}jc')
    jc.set(f'{{{NS}}}val', 'both')
    return p


def load_dc_headers():
    with open(f'{BASE}/dc_section_headers_heb.json', 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    print("Loading data...")
    dc = load_dc()
    pgp = load_pgp()
    chron = load_chron()
    dc_headers = load_dc_headers()

    print("Opening DOCX...")
    doc = Document('C:/Users/chris/Desktop/Hebrew_BOM_MASTER_FINAL_1.docx')

    # ═══ DOCTRINE & COVENANTS ═══
    print("Adding D&C...")

    # Page break to new section (1 column for title)
    add_page_break_1col(doc)
    doc.add_paragraph()  # blank line

    # D&C Title
    add_book_title(doc, 'תורה ובריתות')

    # D&C Introduction with "מבוא" header
    dc_intro = [v for v in dc if v['book'] == 'D&C Intro']
    if dc_intro:
        p = doc.add_heading('מבוא', level=2)
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        etree.SubElement(pPr, f'{{{NS}}}bidi')
        # Add each intro paragraph
        for v in dc_intro:
            add_colophon(doc, v['hebrew'])

    # D&C Sections 1-138
    # sectPr in a paragraph defines the section ENDING at that paragraph
    # Pattern: end prev verses as 2-col → heading+colophon → end as 1-col → verses
    dc_main = [v for v in dc if v['book'] == 'D&C']
    cur_ch = None
    first_section = True
    for v in dc_main:
        ch = v['chapter']
        if ch != cur_ch:
            cur_ch = ch
            # End previous verse section as 2-col (skip for first section)
            if not first_section:
                add_section_break_2col(doc)
            first_section = False
            # Chapter heading + colophon (1-col)
            add_chapter_heading(doc, ch, label='סִימָן')
            header = dc_headers.get(str(ch))
            if header:
                add_colophon(doc, header)
            doc.add_paragraph()  # blank line before two-column verses
            # End heading/colophon section as 1-col
            add_section_break_1col(doc)
        # Verses (in 2-col section, ended by next section break)
        add_verse(doc, v['verse'], v['hebrew'])
    # End final verse section as 2-col
    add_section_break_2col(doc)

    # Official Declarations
    for od_name, od_title in [('OD 1', 'הַכְרָזָה רִשְׁמִית א'), ('OD 2', 'הַכְרָזָה רִשְׁמִית ב')]:
        od_vv = [v for v in dc if v['book'] == od_name]
        if od_vv:
            # Title (1 col) — prev section already ended as 2-col
            add_book_title(doc, od_title)
            add_section_break_1col(doc)  # end title section as 1-col
            # Verses (2 col)
            for v in od_vv:
                add_verse(doc, v['verse'], v['hebrew'])
            add_section_break_2col(doc)  # end verse section as 2-col

    # ═══ PEARL OF GREAT PRICE ═══
    print("Adding PGP...")

    # PGP Title page (1 col)
    add_page_break_1col(doc)
    doc.add_paragraph()
    add_book_title(doc, 'פנינת המחיר הגדול')

    # PGP Introduction (1 col)
    pgp_intro = [v for v in pgp if v['book'] == 'PGP Intro']
    if pgp_intro:
        add_colophon(doc, ' '.join(v['hebrew'] for v in pgp_intro))

    # PGP Books
    pgp_books = ['Moses', 'Abraham', 'JS-Matthew', 'JS-History', 'Articles of Faith']
    pgp_heb = {
        'Moses': 'מֹשֶׁה', 'Abraham': 'אַבְרָהָם',
        'JS-Matthew': 'יוֹסֵף סְמִית — מַתִּתְיָהוּ',
        'JS-History': 'יוֹסֵף סְמִית — הִיסְטוֹרְיָה',
        'Articles of Faith': 'עִקְרֵי הָאֱמוּנָה',
    }
    single_ch = {'JS-Matthew', 'JS-History', 'Articles of Faith'}

    for bk in pgp_books:
        bk_vv = [v for v in pgp if v['book'] == bk]
        if not bk_vv:
            continue

        # Book title (1 col)
        add_book_title(doc, pgp_heb[bk])
        add_section_break_1col(doc)  # end title section as 1-col

        # Verses (2 col)
        cur_ch = None
        for v in bk_vv:
            ch = v['chapter']
            if ch != cur_ch:
                cur_ch = ch
                if bk not in single_ch or ch != 1:
                    add_chapter_heading(doc, ch)
            add_verse(doc, v['verse'], v['hebrew'])
        add_section_break_2col(doc)  # end verse section as 2-col

    # Save
    out = 'C:/Users/chris/Desktop/Hebrew_Triple_Combination.docx'
    print(f"Saving to {out}...")
    doc.save(out)
    print("Done!")


if __name__ == '__main__':
    main()
